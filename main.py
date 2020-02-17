from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import matplotlib.pyplot as plt
import pandas as pd
import os

# Some Inputs
ip_file_name = "LAB-1.cpp"
op_file_name = 'Algo_LAB_AS.docx'

# compiling and running c++ file
os.system("g++ " + ip_file_name)
os.system("./a.out")
colorV = ('red', 'blue', 'green')

# Creating document object
try:
    document = Document(op_file_name)
except:
    document = Document()


def plot_graph(ip):
    count = 0
    for i in ip:
        count += 1
        report_var = 'report' + str(count) + '.csv'
        report = pd.read_csv(report_var)
        plt.plot(report.iloc[:, 0].values, report.iloc[:, 1].values)
    # report1 = pd.read_csv('report1.csv')
    # report2 = pd.read_csv('report2.csv')
    # report3 = pd.read_csv('report3.csv')
    # plt.plot(report1.iloc[:, 0].values, report1.iloc[:, 1].values, color='red')
    # plt.plot(report2.iloc[:, 0].values, report2.iloc[:, 1].values, color='blue')
    # plt.plot(report3.iloc[:, 0].values, report3.iloc[:, 1].values, color='green')

    title_v = ''
    if len(ip) > 1:
        title_v = 'comparing '
        for i in ip:
            if i == ip[0]:
                title_v = title_v + ' ' + i
            else:
                title_v = title_v + ', ' + i
    else:
        title_v = ip[0]

    plt.title(title_v)
    plt.legend(ip)
    plt.xlabel('n')
    plt.ylabel('time taken')
    plt.savefig('fig.png')


# defining top-margin to 0.5inches
sections = document.sections
section = sections[0]
section.top_margin = Inches(0.5)

# Heading like LAB1
paragraph = document.add_paragraph()
# run = paragraph.add_run('LAB {}'.format(lab_no))
run = paragraph.add_run(ip_file_name[:-4])
font = run.font
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
font.size = Pt(22)
font.name = 'Arial'
font.bold = True
# paragraph = document.add_paragraph('LAB1')
# paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
# paragraph.style = document.styles['Heading 1']


# Aim: by aim var
paragraph = document.add_paragraph()
run = paragraph.add_run('Aim: ')
run.font.bold = True
run.font.name = 'sans serif'

# Reading from cpp file
f = open(ip_file_name, 'r')

aim = ''
for line in f:
    if line[0:2] == '/*':
        aim = aim + line[2:-1] + ' '
    elif line[-3:-1] == '*/':
        aim = aim + line[:-3]
        f.close()
        break
    else:
        aim = aim + line[:-1] + ' '
run = paragraph.add_run(aim)
run.font.name = 'sans serif'

# Codes
f = open(ip_file_name, 'r')
codes = ''
k = True

for line in f:
    if k:
        if line[-3:-1] == '*/':
            k = False
    elif k==False and line[-3:-1] == '*/':
        ips = line[2:-3].split()
    else:
        codes = codes + line
f.close()
paragraph = document.add_paragraph()
run = paragraph.add_run('Codes: \n')
run.font.bold = True
run.font.name = 'sans serif'
run = paragraph.add_run(codes)
run.font.name = 'sans serif'
document.add_page_break()

# Output Section
paragraph = document.add_paragraph()
run = paragraph.add_run('Output: ')
run.font.bold = True
run.font.name = 'sans serif'

# Adding graph to docs
plot_graph(ips)
document.add_picture('fig.png', width=Cm(15), height=Cm(11.66))
document.add_page_break()

document.save(op_file_name)

for i in range(len(ips)):
    os.remove('report'+str(i+1)+'.csv')
os.remove('a.out')
